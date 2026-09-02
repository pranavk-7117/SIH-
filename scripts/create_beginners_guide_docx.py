from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "BHUMI-FUSE_Beginners_Guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F8FAFC"
AMBER = "FFF7ED"
GREEN = "F0FDF4"
RED = "FEF2F2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for side in ("top", "bottom", "start", "end"):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), "120" if side in ("start", "end") else "80")
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_cell_margins(table)
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(width)


def style_run(run, bold=False, italic=False, color=INK, size=11) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_para(doc, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        style_run(p.add_run(item))


def add_numbers(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        style_run(p.add_run(item))


def add_heading(doc, text: str, level=1):
    p = doc.add_heading("", level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        size, color = 16, BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size, color = 13, BLUE
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size, color = 12, DARK_BLUE
    style_run(p.add_run(text), bold=True, color=color, size=size)
    return p


def add_callout(doc, title: str, body: str, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(title), bold=True, color=DARK_BLUE, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    style_run(p2.add_run(body), color=INK, size=10)
    doc.add_paragraph()


def add_key_value_table(doc, rows: list[tuple[str, str]], header=("Item", "Details")) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.875, 4.625])
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        style_run(cell.paragraphs[0].add_run(text), bold=True, color=DARK_BLUE)
    for key, value in rows:
        cells = table.add_row().cells
        style_run(cells[0].paragraphs[0].add_run(key), bold=True)
        style_run(cells[1].paragraphs[0].add_run(value))
    doc.add_paragraph()


def add_three_col_table(doc, headers: list[str], rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [1.5, 2.5, 2.5])
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        style_run(cell.paragraphs[0].add_run(text), bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            style_run(cells[i].paragraphs[0].add_run(value), size=10)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header_p.add_run("BHUMI-FUSE Beginner's Guide"), color=MUTED, size=9)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer_p.add_run("Prototype guide | AI assists review, it does not decide land title"), color=MUTED, size=9)


def build() -> None:
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    style_run(title.add_run("BHUMI-FUSE Beginner's Guide"), bold=True, color=RGBColor(11, 37, 69), size=26)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    style_run(subtitle.add_run("A detailed walkthrough of the SIH 2026 land-record harmonization prototype"), color=MUTED, size=12)

    add_callout(
        doc,
        "Core idea",
        "BHUMI-FUSE compares legal-like parcel boundaries with physical evidence, highlights discrepancies, scores confidence transparently, and routes uncertain cases to human review. It never decides land ownership.",
        fill=GREEN,
    )

    add_heading(doc, "1. What Was Built")
    add_para(doc, "This project contains a working judge-facing prototype with a frontend web app, an API layer, deterministic demo data, and scripts for real-data preparation.")
    add_bullets(doc, [
        "Frontend web app: six screens for maps, extraction, evidence graph, harmonization, discrepancy ranking, and review.",
        "Backend API: FastAPI implementation plus a local Node mock API used when Docker is unavailable.",
        "Data scripts: tools to fetch real layers and generate the clearly labelled simulated cadastral layer.",
        "Deployment files: Docker Compose, backend Dockerfile, and frontend Dockerfile for the intended full-stack run.",
    ])

    add_heading(doc, "2. Why It Exists")
    add_para(doc, "Land-record evidence often comes from different sources, dates, accuracies, and authorities. BHUMI-FUSE demonstrates how these layers can be compared without overwriting the original record.")
    add_callout(doc, "Safety guardrail", "The app does not make legal decisions. Low-confidence cases become Needs Review / Do Not Decide, and every reviewer action creates a new versioned record.", fill=AMBER)

    add_heading(doc, "3. Tech Stack")
    add_three_col_table(doc, ["Layer", "Technology", "Purpose"], [
        ("Frontend", "React, TypeScript, Vite", "Builds the interactive user interface and local development server."),
        ("Map UI", "MapLibre GL JS + SVG overlay", "Provides map controls and reliable visible geometry for the demo."),
        ("Styling", "Tailwind CSS + custom CSS", "Creates the dark sidebar, cards, badges, maps, and evidence panels."),
        ("Graph", "react-force-graph-2d", "Renders the interactive Spatial Evidence Graph."),
        ("Backend", "Python, FastAPI, Pydantic", "Defines API endpoints and structured request/response models."),
        ("Geospatial", "GeoPandas, Shapely, Rasterio, PyProj, Pyogrio", "Intended geospatial processing stack for real deployment."),
        ("Analysis", "NetworkX, OpenCV, SciPy", "Evidence graph, boundary extraction, and affine registration."),
        ("Database", "PostgreSQL + PostGIS", "Intended spatial database in Docker Compose."),
    ])

    add_heading(doc, "4. Main Files")
    add_key_value_table(doc, [
        ("docker-compose.yml", "Defines PostGIS, FastAPI, and Vite services for one-command deployment."),
        ("backend/app/main.py", "FastAPI app with ingest, extraction, graph, harmonization, validation, fusion, review, audit, and export endpoints."),
        ("scripts/mock_api_server.mjs", "Local API server currently used to run the demo on this machine without Docker."),
        ("frontend/src/main.tsx", "Main React app containing all six screens and API integrations."),
        ("frontend/src/styles.css", "Visual styling for layout, sidebar, badges, maps, graph, cards, and overlays."),
        ("scripts/fetch_real_layers.py", "Fetches OSM context and writes a manifest for real layer caching."),
        ("scripts/generate_simulated_cadastral.py", "Creates deterministic simulated cadastral boundaries with a fixed seed."),
    ])

    add_heading(doc, "5. The Six Frontend Screens")
    add_heading(doc, "Corpus Overview", level=2)
    add_para(doc, "Shows big KPI cards: parcels processed, conflicts ranked, cases pending review, and versioned actions. It also explains the confidence formula.")
    add_heading(doc, "Source Viewer", level=2)
    add_para(doc, "Shows the core mismatch: amber simulated legal-like parcels, blue real-footprint evidence, gray road context, and purple control points.")
    add_heading(doc, "AI Boundary Extraction", level=2)
    add_para(doc, "Shows extracted boundary observations and confidence values. The prototype presents this as classical OpenCV-style boundary extraction.")
    add_heading(doc, "Spatial Evidence Graph", level=2)
    add_para(doc, "Uses react-force-graph-2d to show parcels, buildings, and control points as nodes. Edges represent evidence relationships and confidence.")
    add_heading(doc, "Harmonized View", level=2)
    add_para(doc, "Shows affine alignment, residual arrows, topology-like validation, and temporal conflict classifications.")
    add_heading(doc, "Conflict Heatmap + Evidence Card", level=2)
    add_para(doc, "Ranks discrepancy cases and lets reviewers Accept, Reject, Adjust, or Escalate. The Evidence Card shows residual distance, confidence, temporal classification, reason, and score breakdown.")

    add_heading(doc, "6. Backend Endpoints")
    add_three_col_table(doc, ["Endpoint", "Role", "Output"], [
        ("GET /health", "Health check", "Service status."),
        ("GET /demo-data", "Demo source data", "GeoJSON layers and bounds."),
        ("POST /ingest", "Layer ingestion", "Versioned ingest record."),
        ("POST /extract", "Boundary extraction", "Boundary observations and confidence."),
        ("GET /graph", "Evidence graph", "Nodes and edges."),
        ("POST /correspond", "Candidate matching", "Ranked parcel-building correspondences."),
        ("POST /harmonize", "Registration", "Affine transform, aligned geometry, residuals."),
        ("POST /validate", "Topology checks", "Accepted or needs-review status."),
        ("POST /temporal-conflict", "Temporal classification", "Registration error, genuine change, or needs review."),
        ("POST /fuse", "Confidence scoring", "Score, state, reason, and breakdown."),
        ("GET /discrepancy-map", "Case ranking", "Ranked discrepancy cases."),
        ("POST /review", "Human decision", "New versioned review record."),
        ("GET /audit/{case_id}", "Audit trail", "History for a case."),
        ("GET /export", "Download", "GeoJSON export."),
    ])

    add_heading(doc, "7. Data and Provenance")
    add_para(doc, "The prototype is designed around real OSM, building-footprint, administrative, and imagery sources. The only intentionally simulated legal-like layer is the cadastral/ownership-style parcel boundary.")
    add_bullets(doc, [
        "Simulated layer: legal-like cadastral boundary, always labelled as not official.",
        "Real or real-intended layers: OSM roads, building footprints, administrative context, current imagery, and historical imagery.",
        "Synthetic control metadata: control points use real-like coordinates with synthetic accuracy values.",
    ])

    add_heading(doc, "8. How Confidence Scoring Works")
    add_para(doc, "The score is rule-based and transparent. This is deliberate: judges and officials can inspect the reason instead of trusting a black-box model.")
    add_key_value_table(doc, [
        ("Authority", "How authoritative the source is considered."),
        ("Positional accuracy", "How spatially reliable the evidence is."),
        ("Temporal relevance", "Whether the timestamp matches the comparison task."),
        ("Completeness", "Whether enough evidence exists."),
        ("Cross-source agreement", "How strongly sources agree spatially."),
        ("Temporal adjustment", "Registration errors are penalized; genuine change is not penalized as bad registration."),
    ], header=("Score part", "Meaning"))

    add_heading(doc, "9. Review and Audit Flow")
    add_para(doc, "Review actions are intentionally versioned. When a reviewer clicks Accept, Reject, Adjust, or Escalate, the app stores a new record with a timestamp. Prior evidence and versions remain intact.")
    add_callout(doc, "Do Not Decide threshold", "If confidence is below the threshold or evidence is mixed, the case is routed to Needs Review / Do Not Decide instead of being force-resolved.", fill=RED)

    add_heading(doc, "10. How to Run")
    add_heading(doc, "Current local demo", level=2)
    add_para(doc, "Because Docker is not installed on this machine, the live demo uses the local Node mock API and Vite frontend.")
    add_key_value_table(doc, [
        ("API", "node scripts/mock_api_server.mjs"),
        ("Frontend", "cd frontend, then npm run dev -- --host 0.0.0.0"),
        ("Open", "http://localhost:5173"),
        ("API health", "http://localhost:8000/health"),
    ])
    add_heading(doc, "Intended Docker demo", level=2)
    add_para(doc, "On a machine with Docker installed, run docker compose up --build from the repository root.")

    add_heading(doc, "11. Suggested Judge Demo Script")
    add_numbers(doc, [
        "Start on Source Viewer and point out the amber simulated legal-like parcels and blue physical footprint evidence.",
        "Explain that the tool does not decide ownership or title.",
        "Open Boundary Extraction and show confidence-based observations.",
        "Open Evidence Graph and click a node to show interlinked evidence.",
        "Open Harmonized View and explain residual arrows.",
        "Point out registration_error, genuine_change, and needs_review classifications.",
        "Open Discrepancy Map, click the top case, and show the Evidence Card.",
        "Click Escalate or Adjust and explain that a new version was created without overwriting the source.",
    ])

    add_heading(doc, "12. Prototype-Level Simplifications")
    add_bullets(doc, [
        "The currently running demo uses deterministic geometry instead of fully downloaded imagery.",
        "The real-data fetch script is scaffolded, but not every real source is fully automated yet.",
        "PostGIS is included in Docker Compose, but the local mock API does not use it.",
        "Registration and boundary extraction are represented with deterministic demo outputs.",
        "The Node mock API exists only so the demo can run on this machine without Docker.",
    ])

    add_heading(doc, "Final Summary")
    add_para(doc, "BHUMI-FUSE is an evidence-fusion and discrepancy-prioritization prototype. It helps officials inspect legal-vs-physical land-record conflicts, understand why a case is risky, and preserve a reviewable audit trail.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
